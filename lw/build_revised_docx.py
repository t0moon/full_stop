import re
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "aigc_brand_image_report_revised.md"
TARGET = ROOT / "aigc_brand_image_report_revised.docx"
PUBLIC_TARGET = ROOT / "多模态话语视角下AIGC视频广告品牌形象建构策略研究_修订版.docx"


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name, east_asia, latin, size in [
        ("Heading 1", "黑体", "SimHei", 14),
        ("Heading 2", "黑体", "SimHei", 12),
    ]:
        style = doc.styles[style_name]
        style.font.name = latin
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    set_run_font(run, "宋体", 12)


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    set_run_font(run, "黑体", 14 if level == 1 else 12, bold=True)


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    set_run_font(run, "黑体", 16, bold=True)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    set_run_font(run, "黑体" if bold else "宋体", 10.5, bold=bold)


def add_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2 or not is_table_separator(table_lines[1]):
        for line in table_lines:
            add_paragraph(doc, line)
        return

    header = parse_table_row(table_lines[0])
    data_rows = [parse_table_row(line) for line in table_lines[2:]]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    for index, text in enumerate(header):
        set_cell_text(table.rows[0].cells[index], text, bold=True)

    for values in data_rows:
        row = table.add_row()
        padded = values + [""] * (len(header) - len(values))
        for index, text in enumerate(padded[: len(header)]):
            set_cell_text(row.cells[index], text)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    section.start_type = WD_SECTION.CONTINUOUS

    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue

        if is_table_line(line):
            table_lines = [line]
            index += 1
            while index < len(lines) and is_table_line(lines[index].strip()):
                table_lines.append(lines[index].strip())
                index += 1
            add_table(doc, table_lines)
            continue

        if line.startswith("# "):
            add_title(doc, line[2:].strip())
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        else:
            add_paragraph(doc, line)
        index += 1

    doc.save(TARGET)
    copy2(TARGET, PUBLIC_TARGET)


if __name__ == "__main__":
    build_docx()
