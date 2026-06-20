from __future__ import annotations

from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "成稿01.docx"

TITLE = "多模态话语视角下AIGC视频广告品牌形象建构策略研究"
BLANK = "____________________"


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    none = {"val": "nil"}
    set_cell_border(
        cell,
        left=none,
        top=none,
        right=none,
        bottom=none,
        insideH=none,
        insideV=none,
    )


def set_paragraph_format(para, *, first_line_indent: float = 0, line: float = 20) -> None:
    para.paragraph_format.first_line_indent = Pt(first_line_indent)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(line)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def add_cell_paragraph(cell, text: str, *, bold: bool = False, indent: float = 0, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    para.alignment = align
    set_paragraph_format(para, first_line_indent=indent)
    run = para.add_run(text)
    set_run_font(run, "宋体", 10.5, bold=bold)


SECTIONS: list[tuple[str, list[tuple[str, str]]]] = [
    (
        "一、课题来源及选题依据",
        [
            ("head2", "（一）课题来源"),
            ("head3", "现实背景"),
            (
                "body",
                "近两年，生成式人工智能技术快速进入广告内容生产链条，应用范围已从文案撰写、图片生成扩展到脚本构思、分镜设计、角色塑造、配音配乐、视频生成与智能剪辑等多个环节。AIGC对广告行业的影响已不只是单点提效，而是对内容生产方式、创意组织方式和传播表达方式的整体改造。",
            ),
            (
                "body",
                "在这一变化中，视频广告尤为值得关注。相较于平面广告和纯文本广告，视频广告具有语言、画面、音乐、音效、字幕、镜头运动和剪辑节奏等多种符号资源协同运作的特征，是一种典型的动态多模态语篇。品牌越来越依赖视频广告来传递品牌理念、产品价值、生活方式想象与情感氛围，因此，AIGC进入视频广告创作后，其对品牌传播的影响具有较强的研究价值。",
            ),
            ("head3", "行业发展现状"),
            (
                "body",
                "现有行业实践表明，AIGC能够显著降低视频广告的制作门槛，并强化视觉奇观、场景转换和风格模拟能力。AIGC已经从创意辅助工具逐渐进入广告生产主流程，品牌方、平台和内容团队均在尝试将其用于脚本生成、角色塑造、画面合成与后期制作等环节。",
            ),
            (
                "body",
                "但与此同时，AIGC品牌视频广告也经常面临品牌识别度不足、情感表达单薄、叙事逻辑松散、视觉新奇与品牌意义脱节、平台风格趋同等问题。这说明，AIGC并非只是替代传统制作流程的工具，它还改变了广告中不同模态资源的组织方式，进而影响品牌形象的呈现效果。",
            ),
            ("head3", "研究问题提出"),
            (
                "body",
                "因此，本课题的问题意识不再停留于“AIGC能否用于广告制作”，也不再将重点放在人工广告与AIGC广告的优劣比较上，而是聚焦于一个更具解释力的问题：在AIGC品牌视频广告这一单一研究对象中，语言、视觉、听觉及其协同关系是如何共同呈现品牌形象的？又形成了哪些相对稳定的呈现模式、常见问题与优化方向？",
            ),
            (
                "body",
                "基于此，本研究将现实中的“AIGC品牌视频广告兴起”转化为学术上的“多模态品牌形象呈现问题”。研究不以开发视频生成技术或评价工具效率为目标，而是把AIGC品牌视频广告视为一种新型动态多模态文本，分析其如何调动文字、图像、色彩、音乐、音效、镜头与节奏等资源来呈现品牌意义，并进一步总结其相对稳定的呈现模式。",
            ),
            ("head2", "（二）理论依据"),
            ("head3", "1. 理论一"),
            (
                "body",
                "理论来源：理论一主要指多模态话语分析理论，具体以张德禄多模态话语分析综合框架为总体框架，并吸收韩礼德三大元功能理论、Kress和van Leeuwen视觉语法以及模态协同理论的相关思想。",
            ),
            (
                "body",
                "核心观点：该理论认为意义生产并不依赖单一语言符号，而是由语言、视觉、听觉等多种模态共同完成。张德禄框架强调从语境层、内容层和表达层理解多模态话语的生成过程，并关注模态之间的互补、强化、协调与弱关联关系。",
            ),
            (
                "body",
                "与研究对象的适配：AIGC品牌视频广告本身是动态多模态文本，天然包含文案、画面、音乐、音效、镜头与节奏等资源。以该理论为基础，有助于解释不同模态如何共同呈现品牌意义，并将传统多模态分析中的“内容层”进一步转化为更贴近本研究对象的“品牌形象呈现层”。",
            ),
            ("head3", "2. 理论二"),
            (
                "body",
                "理论来源：理论二主要指品牌形象理论，综合品牌意义建构研究与品牌形象分类研究，重点参考Keller关于品牌联想的界定，以及Park、Jaworski与MacInnis关于品牌形象类型的讨论。",
            ),
            (
                "body",
                "核心观点：品牌形象并不等同于品牌名称、标识或产品露出的简单呈现，而是由功能认知、情感联结、符号认同和价值表达等维度共同构成的意义结构。广告中的品牌形象需要依靠多种传播线索共同支撑，才能形成相对稳定且可识别的品牌联想。",
            ),
            (
                "body",
                "与研究对象的适配：本文不直接考察消费者心智中的最终效果，而是把品牌形象作为广告文本分析维度，用于判断AIGC品牌视频广告重点呈现了何种品牌意义。该理论能够为后续编码、比较和模式归纳提供明确的判断标准。",
            ),
        ],
    ),
    (
        "二、文献综述",
        [
            ("head2", "（一）国内外研究现状"),
            ("head3", "1. 研究主题A"),
            (
                "body",
                "研究主题A主要对应多模态话语分析理论及其在动态视听文本中的应用研究。现有研究表明，多模态话语分析已经形成较成熟的“总框架+分项工具”组合模式，即以综合框架界定分析层级，再调用元功能理论、视觉语法和模态协同理论解释不同模态的意义承担方式。相关研究已广泛用于纪录片、短视频、外宣语篇、文化节目和跨文化传播视频，证明该方法具备分析动态视听文本和复杂意义建构的能力。",
            ),
            ("head3", "2. 研究主题B"),
            (
                "body",
                "研究主题B主要对应视频广告与品牌形象建构研究。已有研究通常围绕广告诉求、叙事结构、人物设置、品牌呈现、情感表达和视听风格等维度分析视频广告，并将品牌形象理解为由功能认知、情感联结、符号认同和价值表达等多层意义共同构成的结果。视频广告研究为本文提供了文本分析的维度基础，品牌形象研究则为本文提供了意义判断的核心维度。",
            ),
            ("head3", "3. 研究主题C"),
            (
                "body",
                "研究主题C主要对应AIGC视频广告研究。现有成果大体集中在三类问题：一是AIGC如何进入广告生产流程，二是AIGC广告与人工广告在传播效果和受众反应上的差异，三是从业者如何评价和采用AI视频广告工具。总体来看，AIGC视频广告研究更多关注技术应用、广告效果和工具采纳，而对广告文本内部如何组织意义、如何呈现品牌形象、如何形成相对稳定的表达策略，讨论仍然不足。",
            ),
            ("head2", "（二）文献评述"),
            ("head3", "已有研究成果"),
            (
                "body",
                "综合以上文献，现有研究已经形成三条彼此关联的线索：一是多模态话语分析已经形成较成熟的理论基础和应用路径；二是视频广告研究和品牌形象研究已经为广告文本分析提供了较稳定的内容维度和意义维度；三是AIGC广告研究已经扩展到技术应用、广告效果、AI披露、工具采纳和内容质量评价等方面。",
            ),
            ("head3", "已有研究不足"),
            (
                "body",
                "但现有研究仍存在明显缺口。其一，多模态话语分析的应用对象仍以公共传播文本为主，对AIGC品牌视频广告这一商业传播文本关注不足。其二，视频广告研究与多模态话语分析之间的结合还不够紧密。其三，品牌形象研究较少进入AIGC语境，也较少从语言、视觉、听觉及其协同关系的角度讨论品牌形象呈现与建构方式。其四，AIGC广告研究虽然开始关注真实性、创造性、可信度和工具采纳，但对广告文本内部如何组织品牌意义、如何形成稳定表达策略，仍缺乏系统分析。",
            ),
            ("head3", "本研究切入点"),
            (
                "body",
                "基于此，本文将以AIGC品牌视频广告为单一研究对象，从多模态话语视角考察广告文本中的品牌形象呈现方式，重点分析语言模态、视觉模态、听觉模态及其协同关系如何共同塑造品牌功能形象、情感形象、符号形象和价值形象，并在跨样本比较基础上总结其相对稳定的呈现模式、常见问题和优化方向。",
            ),
        ],
    ),
    (
        "三、研究内容、研究目的与研究意义",
        [
            ("head2", "（一）研究内容"),
            ("head3", "研究对象"),
            (
                "body",
                "本研究的对象明确限定为AIGC品牌视频广告这一单一对象，不再设置“人工广告”与“AIGC广告”的双组比较。研究聚焦的问题是AIGC品牌视频广告如何在广告文本内部呈现品牌形象，而不是比较不同生产方式的优劣。",
            ),
            ("head3", "研究范围"),
            (
                "body",
                "在总体界定上，本研究所指的总体是2024年1月1日至2026年6月1日期间公开发布、且AIGC参与创作环节可被核验的品牌视频广告。样本拟采用目的性抽样与热度筛选相结合的方式，从品牌官方发布渠道、主流AIGC平台公开案例库以及权威行业报道中建立候选池，最终选取60条AIGC品牌视频广告作为正式研究样本。",
            ),
            (
                "body",
                "研究范围还包括样本纳入与排除边界的限定。纳入样本需具有明确品牌传播属性、AIGC介入环节可核验、视频文本完整且适合切分与编码；无法核验AIGC介入方式、非商业属性视频以及素材残缺作品不纳入正式样本。",
            ),
            ("head3", "研究框架"),
            (
                "body",
                "本文在多模态话语分析综合框架基础上，构建适用于AIGC品牌视频广告研究的四层分析框架，即语境层、表达层、品牌形象呈现层和模态关系层。语境层用于识别广告所处的传播场景、诉求导向和品牌传播重心；表达层用于编码语言、视觉和听觉资源的具体使用方式；品牌形象呈现层用于判断广告文本重点建构的是何种品牌意义；模态关系层则用于分析语言、视觉、听觉之间是强化、互补、协调还是弱关联，从而解释品牌意义如何被整体建构。",
            ),
            ("head3", "章节安排"),
            (
                "body",
                "全文拟由绪论、五章正文和结语构成。现有材料已经明确的主体结构包括：文献综述与理论基础、研究设计、AIGC品牌广告的话语情境特征、AIGC品牌广告的表达层特征，以及围绕话语机制与品牌形象呈现展开的分析部分。各章节共同服务于“广告文本中的品牌形象如何被多模态呈现和组织”这一核心问题。",
            ),
            ("head2", "（二）研究目的"),
            ("head3", "理论目的"),
            (
                "body",
                "本研究的理论目的，一是构建适用于AIGC品牌视频广告研究的量化编码框架，回答这一类广告在多模态资源调动上“呈现了什么”以及“如何呈现”的问题；二是通过统计分析归纳AIGC品牌视频广告文本中品牌形象呈现的主要分布特征和组合规律，进一步提炼出若干具有解释力的呈现模式。",
            ),
            ("head3", "现实目的"),
            (
                "body",
                "本研究的现实目的，是在量化结果和典型案例解释的基础上，为品牌方和创作团队提出更有针对性的优化思路，使AIGC广告创作更好地服务于品牌传播，而不是止步于视觉展示和技术奇观。",
            ),
            ("head2", "（三）研究意义"),
            ("head3", "理论意义"),
            (
                "body",
                "在理论层面，研究有助于推动AIGC广告研究由“工具应用”与“传播效果”转向“意义生产”与“广告文本中的品牌形象呈现”，丰富现有研究议题；同时能够将品牌形象研究与多模态话语分析结合起来，为AIGC语境下的视频广告研究提供新的分析框架，并拓展广告学与传播学在人工智能内容生产情境中的研究边界。",
            ),
            ("head3", "实践意义"),
            (
                "body",
                "在实践层面，研究能够为品牌方理解AIGC视频广告的优势与局限提供参考，避免创作中过度依赖技术奇观而忽视品牌一致性；也可为AIGC广告创作团队提供更明确的品牌表达思路，提升人机协同创作的有效性，并推动相关研究从零散观察走向相对系统的经验分析。",
            ),
        ],
    ),
    (
        "四、研究方法与研究设计",
        [
            ("head2", "（一）研究方法"),
            ("head3", "内容分析法"),
            (
                "body",
                "本研究以AIGC品牌视频广告为研究文本，对其多模态资源进行系统编码和分类，通过量化方式考察广告文本中的稳定特征。由于研究关注的是广告文本中的品牌形象如何被呈现，内容分析法能够较好实现从文本观察到统计归纳的转换，并支持不同样本之间的跨案例比较。",
            ),
            ("head3", "深度访谈法"),
            ("blank", ""),
            ("head3", "参与观察法"),
            ("blank", ""),
            ("head3", "个案研究法"),
            (
                "body",
                "在量化结果基础上，本研究将选取具有代表性的个案对统计结果进行补充解释，避免仅有数字而缺乏文本意义分析。通过典型案例补充分析，可以形成“量化编码为主、定性解释为辅”的研究路径。",
            ),
            ("head3", "多模态分析法"),
            (
                "body",
                "本研究将AIGC品牌视频广告视为动态多模态文本，从语言、视觉、听觉及其协同关系入手展开分析。具体操作上，研究拟借助视频标注工具对样本进行转录、切分和分层标注，将字幕、画面、声音和镜头等信息整理为可统计的编码材料，并在四层分析框架下识别不同模态资源的分工与协同方式。",
            ),
            ("head2", "（二）研究设计"),
            ("head3", "样本来源"),
            (
                "body",
                "本研究所需样本主要来自品牌官方账号、品牌传播平台、主流AIGC视频工具公开案例库以及行业报道中可核验的品牌广告作品。相关视频文本公开可得，便于重复观看、转录、切分和编码，具备开展量化内容分析的基本条件。",
            ),
            ("head3", "样本筛选"),
            (
                "body",
                "样本采用目的性抽样与热度筛选相结合的方式建立候选池。纳入标准包括：具有明确品牌传播属性；AIGC参与创作的关键环节可通过品牌说明、创作者披露、平台案例说明或权威报道获得核验；视频文本完整，可反复观看并适合进行镜头切分、字幕转录和多模态编码；样本尽可能覆盖不同品牌品类、传播平台和表达风格。排除标准包括：无法核验AIGC介入方式的作品、非商业属性视频以及素材残缺或无法完整获取的视频。",
            ),
            ("head3", "数据收集"),
            (
                "body",
                "在数据收集阶段，研究将以“单支广告视频”为样本单位和最终统计单位，以“镜头/场景片段”及其对应的字幕、配音、音乐与音效片段为编码参照单位。研究拟借助视频抽帧、字幕提取、OCR识别和视频标注工具辅助整理原始材料，对每支视频进行镜头切分与分层标注，再将片段层观察结果统一汇总为视频层变量。",
            ),
            ("head3", "数据分析流程"),
            (
                "body",
                "本研究的数据分析流程为“样本库构建—编码维度确定—预编码与正式编码—统计分析—典型案例解释—模式归纳与问题诊断”。正式分析时，将对编码结果进行频数统计、占比分析和交叉比较，呈现AIGC品牌视频广告在模态使用和品牌形象呈现方面的分布规律，并据此归纳不同的呈现模式。为保证研究质量，还将通过明确工作定义、预编码、双人复编码和文本回核等方式控制信度与解释边界。",
            ),
        ],
    ),
    (
        "五、论文框架（重点）",
        [
            ("head2", "绪论"),
            (
                "body",
                "绪论部分主要交代研究背景与问题提出、研究目的与研究意义、研究内容与研究问题、研究思路与研究方法，以及创新点与整体研究框架。",
            ),
            ("head3", "第一章"),
            (
                "body",
                "第一章拟为“文献综述与理论基础”，主要梳理多模态话语分析理论研究、视频广告中的品牌形象建构研究、AIGC视频广告研究，并在此基础上形成文献评述与研究切入点。",
            ),
            ("head3", "第二章"),
            (
                "body",
                "第二章拟为“研究设计”，主要说明研究对象与样本来源、理论框架构建、编码指标设置、数据收集与编码流程，以及信度与效度控制。",
            ),
            ("head3", "第三章"),
            (
                "body",
                "第三章拟分析AIGC品牌广告的话语情境特征，重点讨论AIGC时代的品牌传播生态、广告所处的文化语境与情境语境，以及不同传播场景中的品牌意义生产机制。",
            ),
            ("head3", "第四章"),
            (
                "body",
                "第四章拟分析AIGC品牌广告的表达层特征，分别讨论语言模态、视觉模态、听觉模态及其组合方式，呈现广告文本在多模态资源调动上的主要特征。",
            ),
            ("head3", "第五章"),
            (
                "body",
                "第五章拟在前文分析基础上讨论AIGC品牌广告的话语机制与品牌形象呈现逻辑，重点归纳不同品牌意义的建构方式、相对稳定的呈现模式及其存在的问题。",
            ),
            ("head2", "结语"),
            ("blank", ""),
        ],
    ),
    (
        "六、已有研究基础",
        [
            ("head2", "（一）已有工作"),
            ("head3", "文献阅读"),
            (
                "body",
                "现阶段，本课题已围绕多模态话语分析理论与应用、视频广告与品牌形象建构、AIGC视频广告等方向完成初步文献梳理，基本形成“多模态话语分析理论与应用—视频广告与品牌形象建构—AIGC视频广告—研究缺口”的综述框架。",
            ),
            ("head3", "数据收集"),
            (
                "body",
                "在材料准备方面，研究已经完成选题聚焦，明确将研究对象限定为AIGC品牌视频广告，并开始整理样本候选池，初步提炼出适用于本研究的样本筛选原则与案例核验要求。",
            ),
            ("head3", "理论梳理"),
            (
                "body",
                "在理论准备方面，研究已结合张德禄多模态话语分析综合框架、视觉语法、模态协同理论以及品牌形象相关研究，初步形成适用于本研究的分析层级和解释思路，并开始将其转化为后续编码设计所需的操作性框架。",
            ),
            ("head2", "（二）存在问题"),
            ("head3", "理论问题"),
            (
                "body",
                "目前关于AIGC广告应用与效果的文献相对较多，但与品牌形象建构直接相关的经典研究和可用于细化分析维度的成果仍需继续补充，以增强理论支撑的完整性。",
            ),
            ("head3", "方法问题"),
            (
                "body",
                "量化编码有助于提高研究的系统性，但若编码维度界定不清、不同编码者判断标准不一致，容易影响结果稳定性；同时，若过度依赖数字分布，也可能削弱对广告文本细节和品牌意义的解释力。",
            ),
            ("head3", "材料问题"),
            (
                "body",
                "当前AIGC品牌视频广告大多属于人机协同产物，部分作品虽使用AI，但未必适合纳入本研究；部分作品虽具有品牌属性，但AIGC介入程度难以核验。同时，部分变量如“情感形象突出程度”“模态协调度”等仍需进一步明确工作定义，才能进入正式编码。",
            ),
            ("head2", "（三）解决方案"),
            (
                "body",
                "后续将围绕品牌形象维度、视频广告叙事、品牌符号表达与多模态传播等主题继续补充文献，并在已有AIGC研究基础上形成更完整的综述结构。",
            ),
            (
                "body",
                "在样本与方法方面，将对每个候选案例记录品牌属性、发布时间、传播平台、AIGC可核验说明和文本完整度，对界定模糊的案例采取从严处理；同时在正式编码前对各变量给出操作性定义、赋值规则和判断示例，降低编码歧义。",
            ),
            (
                "body",
                "在质量控制方面，将先进行预编码，根据问题调整变量设置；正式编码阶段对部分样本实施双人复编码并进行一致性检验。对于量化结果无法充分解释的部分，再回到典型案例作补充分析，并在涉及技术机制解释时采用审慎措辞，避免把统计分布直接写成确定因果。",
            ),
        ],
    ),
]

SCHEDULE = [
    ("6月", "文献综述"),
    ("6月", "开题报告"),
    ("6月", "开题答辩"),
    ("10月", "初稿"),
    ("1月", "定稿"),
    ("4月", "答辩"),
]


def add_cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, line=35)
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("湘   潭  大  学")
    set_run_font(run, "宋体", 36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, line=35)
    run = p.add_run("硕士研究生学位论文开题报告及工作计划书")
    set_run_font(run, "宋体", 18)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p, line=28)
    run = p.add_run("论文题目：")
    set_run_font(run, "宋体", 16)
    run = p.add_run(TITLE)
    set_run_font(run, "宋体", 16, bold=True)

    info = [
        ("学   号", BLANK),
        ("研究生姓名", BLANK),
        ("所在学院", BLANK),
        ("专   业", BLANK),
        ("研究方向", BLANK),
        ("导师姓名", BLANK),
        ("入学年月", BLANK),
        ("开题时间", BLANK),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in info:
        row = table.add_row()
        row.height = Cm(1.1)
        left = row.cells[0]
        right = row.cells[1]
        left.width = Cm(4.5)
        right.width = Cm(8.5)
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        clear_cell_borders(left)
        clear_cell_borders(right)
        para = left.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(para, line=12)
        run = para.add_run(label)
        set_run_font(run, "宋体", 16)
        para = right.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(para, line=12)
        run = para.add_run(value)
        set_run_font(run, "宋体", 16)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_outer_row_content(cell, section_title: str, items: list[tuple[str, str]]) -> None:
    title_para = cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(title_para, line=20)
    run = title_para.add_run(section_title)
    set_run_font(run, "宋体", 10.5, bold=True)

    for kind, text in items:
        if kind == "head2":
            add_cell_paragraph(cell, text, bold=True)
        elif kind == "head3":
            add_cell_paragraph(cell, text, bold=True)
        elif kind == "body":
            add_cell_paragraph(cell, text, indent=21)
        elif kind == "blank":
            add_cell_paragraph(cell, " ")
        else:
            raise ValueError(kind)


def add_schedule_row(cell) -> None:
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(para, line=20)
    run = para.add_run("七、进度安排")
    set_run_font(run, "宋体", 10.5, bold=True)

    nested = cell.add_table(rows=1, cols=2)
    nested.style = "Table Grid"
    nested.alignment = WD_TABLE_ALIGNMENT.CENTER
    nested.autofit = False
    headers = ["时间", "工作内容"]
    for idx, text in enumerate(headers):
        hdr = nested.rows[0].cells[idx]
        hdr.width = Cm(3.5 if idx == 0 else 10.5)
        hdr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = hdr.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(para, line=20)
        run = para.add_run(text)
        set_run_font(run, "宋体", 10.5, bold=True)
    for left_text, right_text in SCHEDULE:
        row = nested.add_row()
        values = [left_text, right_text]
        for idx, text in enumerate(values):
            c = row.cells[idx]
            c.width = Cm(3.5 if idx == 0 else 10.5)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = c.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(para, line=20)
            run = para.add_run(text)
            set_run_font(run, "宋体", 10.5)


def add_signature_row(cell, title: str, sign_text: str) -> None:
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(para, line=20)
    run = para.add_run(title)
    set_run_font(run, "宋体", 10.5, bold=True)
    add_cell_paragraph(cell, " ")
    add_cell_paragraph(cell, sign_text)


def build_form(doc: Document) -> None:
    row_count = len(SECTIONS) + 1 + 3
    table = doc.add_table(rows=row_count, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4), Cm(4), Cm(4), Cm(4)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for idx, (section_title, items) in enumerate(SECTIONS):
        merged = table.cell(idx, 0).merge(table.cell(idx, 3))
        add_outer_row_content(merged, section_title, items)

    schedule_idx = len(SECTIONS)
    merged = table.cell(schedule_idx, 0).merge(table.cell(schedule_idx, 3))
    add_schedule_row(merged)

    signatures = [
        ("导师审查意见", "导师签字：____________________          年     月     日"),
        ("学位点审查意见", "学位点负责人签字：____________________          年     月     日"),
        ("学院审查意见", "学院负责人签字：____________________          年     月     日"),
    ]
    for offset, (title, sign) in enumerate(signatures, start=1):
        merged = table.cell(schedule_idx + offset, 0).merge(table.cell(schedule_idx + offset, 3))
        add_signature_row(merged, title, sign)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    section.start_type = WD_SECTION_START.NEW_PAGE


def build_doc() -> Path:
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)

    backup = ROOT / f"成稿01_before_formatted_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    copy2(TARGET, backup)

    doc = Document()
    configure_page(doc)
    add_cover_page(doc)
    build_form(doc)
    doc.save(TARGET)
    return backup


if __name__ == "__main__":
    print(build_doc())
