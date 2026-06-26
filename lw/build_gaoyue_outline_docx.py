from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "高月硕士毕业论文开题报告v1.0_修订目录.md"
TARGET = ROOT / "高月硕士毕业论文开题报告v1.0_修订目录.docx"


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = para.add_run(text)
    set_run_font(run, "黑体", 16, bold=True)


def add_outline_line(doc: Document, text: str, level: int) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 14
        bold = True
        left_indent = 0
    elif level == 2:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 14
        bold = True
        left_indent = 0
    elif level == 3:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 12
        bold = True
        left_indent = 18
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 12
        bold = False
        left_indent = 36

    para.paragraph_format.left_indent = Pt(left_indent)
    run = para.add_run(text)
    set_run_font(run, "黑体" if bold else "宋体", size, bold=bold)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if line.startswith("# "):
            add_title(doc, line[2:].strip())
            continue

        for marker, level in [("## ", 1), ("### ", 2), ("#### ", 3), ("##### ", 4)]:
            if line.startswith(marker):
                add_outline_line(doc, line[len(marker) :].strip(), level)
                break

    doc.save(TARGET)


if __name__ == "__main__":
    build_docx()
